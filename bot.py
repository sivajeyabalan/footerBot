import os
import re
import logging
import asyncio
import signal
import sys
from typing import Tuple, Optional, Dict
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from telegram.error import TimedOut, NetworkError, RetryAfter, Conflict
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import subprocess

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Constants
SUPPORTED_EXTENSIONS = {'.docx'}  # Only allow DOCX files
MAX_RETRIES = 3
RETRY_DELAY = 1  # seconds

# Environment variables
PORT = int(os.getenv('PORT', '8443'))
ENVIRONMENT = os.getenv('ENVIRONMENT', 'development')  # 'development' or 'production'
WEBHOOK_URL = os.getenv('WEBHOOK_URL', '')  # Required for production

# Conversation states
WAITING_FOR_NAME, WAITING_FOR_ROLLNO = range(2)

# Store user data temporarily
user_data: Dict[int, Dict] = {}

# Global application instance
application = None

async def cleanup():
    """Clean up resources and temporary files."""
    global application
    try:
        if application:
            logger.info("Stopping application...")
            await application.stop()
            await application.shutdown()
        
        # Clean up temporary files
        for file in os.listdir():
            if file.endswith(('.docx', '.pdf')) and not file == 'requirements.txt':
                try:
                    os.remove(file)
                    logger.info(f"Cleaned up temporary file: {file}")
                except Exception as e:
                    logger.error(f"Error cleaning up file {file}: {e}")
        
        logger.info("Cleanup completed")
    except Exception as e:
        logger.error(f"Error during cleanup: {e}")

def signal_handler(signum, frame):
    """Handle system signals for graceful shutdown."""
    logger.info(f"Received signal {signum}")
    asyncio.create_task(cleanup())
    sys.exit(0)

# Register signal handlers
signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)

async def send_message_with_retry(update: Update, text: str, max_retries: int = MAX_RETRIES) -> bool:
    """Send a message with retry logic."""
    for attempt in range(max_retries):
        try:
            await update.message.reply_text(text)
            return True
        except (TimedOut, NetworkError) as e:
            if attempt < max_retries - 1:
                logger.warning(f"Attempt {attempt + 1} failed: {str(e)}. Retrying...")
                await asyncio.sleep(RETRY_DELAY)
            else:
                logger.error(f"Failed to send message after {max_retries} attempts: {str(e)}")
                return False
        except RetryAfter as e:
            logger.warning(f"Rate limited. Waiting {e.retry_after} seconds...")
            await asyncio.sleep(e.retry_after)
        except Exception as e:
            logger.error(f"Unexpected error while sending message: {str(e)}")
            return False
    return False

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a message when the command /start is issued."""
    welcome_message = (
        "👋 Welcome to the Document Footer Bot!\n\n"
        "To use this bot:\n"
        "1. Send a DOCX file\n"
        "2. When prompted, send your name\n"
        "3. When prompted, send your roll number\n\n"
        "The bot will add this information as a footer to your document, convert it to PDF, and send it back."
    )
    await send_message_with_retry(update, welcome_message)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle incoming documents and start the conversation."""
    try:
        # Get the document
        document = update.message.document
        file_name = document.file_name
        file_ext = os.path.splitext(file_name)[1].lower()
        
        # Get username or user ID as fallback
        user = update.effective_user
        user_identifier = user.username if user.username else str(user.id)
        
        # Check file type
        if file_ext not in SUPPORTED_EXTENSIONS:
            await send_message_with_retry(update, f"❌ Unsupported file type. Please send a DOCX file.")
            return ConversationHandler.END
        
        # Download the file
        file = await context.bot.get_file(document.file_id)
        file_path = f"{user_identifier}_{file_name}"
        await file.download_to_drive(file_path)
        
        # Store file information in user_data
        user_id = update.effective_user.id
        user_data[user_id] = {
            'file_path': file_path,
            'file_ext': file_ext,
            'file_name': file_name,
            'user_identifier': user_identifier
        }
        
        # Ask for name
        await send_message_with_retry(update, "Please send your name:")
        return WAITING_FOR_NAME
        
    except Exception as e:
        logger.error(f"Error handling document: {str(e)}")
        await send_message_with_retry(update, "❌ An error occurred while processing your document. Please try again.")
        return ConversationHandler.END

async def handle_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle the user's name and ask for roll number."""
    user_id = update.effective_user.id
    name = update.message.text.strip()
    
    if not name:
        await send_message_with_retry(update, "❌ Please provide a valid name:")
        return WAITING_FOR_NAME
    
    # Store name in user_data
    user_data[user_id]['name'] = name
    
    # Ask for roll number
    await send_message_with_retry(update, "Please send your roll number:")
    return WAITING_FOR_ROLLNO

async def handle_rollno(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle the user's roll number and process the document."""
    user_id = update.effective_user.id
    roll_no = update.message.text.strip()
    
    if not roll_no:
        await send_message_with_retry(update, "❌ Please provide a valid roll number:")
        return WAITING_FOR_ROLLNO
    
    try:
        # Acknowledge the user that processing is starting
        await send_message_with_retry(update, "🔄 Processing your document. Please wait...")
        
        # Get stored data
        data = user_data[user_id]
        file_path = data['file_path']
        file_ext = data['file_ext']
        name = data['name']
        user_identifier = data['user_identifier']
        
        # Process the document (only DOCX now)
        docx_output_path = await process_docx(file_path, name, roll_no, user_identifier)
        
        # Convert DOCX to PDF
        pdf_output_path = await convert_docx_to_pdf(docx_output_path)
        
        # Send the processed PDF file
        await update.message.reply_document(
            document=open(pdf_output_path, 'rb'),
            filename=os.path.basename(pdf_output_path)
        )
        
        # Clean up
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(docx_output_path):
            os.remove(docx_output_path)
        if os.path.exists(pdf_output_path):
            os.remove(pdf_output_path)
        del user_data[user_id]
        
        await send_message_with_retry(update, "✅ Document processed successfully!")
        return ConversationHandler.END
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        await send_message_with_retry(update, "❌ An error occurred while processing your document. Please try again.")
        # Clean up
        if user_id in user_data:
            if os.path.exists(user_data[user_id]['file_path']):
                os.remove(user_data[user_id]['file_path'])
            del user_data[user_id]
        return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Cancel the conversation and clean up."""
    user_id = update.effective_user.id
    if user_id in user_data:
        if os.path.exists(user_data[user_id]['file_path']):
            os.remove(user_data[user_id]['file_path'])
        del user_data[user_id]
    
    await send_message_with_retry(update, "Operation cancelled. Send /start to begin again.")
    return ConversationHandler.END

async def process_docx(file_path: str, name: str, roll_no: str, user_identifier: str) -> str:
    """Process DOCX file and add footer."""
    doc = docx.Document(file_path)
    
    # Add footer to each section
    for section in doc.sections:
        # Reduce the footer distance to bring it closer to the border
        section.footer_distance = Inches(0.3)
        footer = section.footer
        
        # Clear any existing paragraphs
        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Create a table for the footer with specified width
        table = footer.add_table(rows=1, cols=3, width=Inches(7.5))
        table.autofit = False
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(2.5)
        table.columns[2].width = Inches(2.5)
        
        # Add name (left)
        cell = table.cell(0, 0)
        name_run = cell.paragraphs[0].add_run(f"Name: {name}")
        name_run.font.size = Pt(10)
        name_run.font.name = 'Times New Roman'
        
        # Add roll number (middle)
        cell = table.cell(0, 1)
        roll_run = cell.paragraphs[0].add_run(f"Roll No: {roll_no}")
        roll_run.font.size = Pt(10)
        roll_run.font.name = 'Times New Roman'
        
        # Add page number (right)
        cell = table.cell(0, 2)
        page_run = cell.paragraphs[0].add_run("Page no:")
        page_run.font.size = Pt(10)
        page_run.font.name = 'Times New Roman'
        
        # Add page number field using proper XML namespace
        page_number = cell.paragraphs[0].add_run()
        page_number.font.size = Pt(10)
        page_number.font.name = 'Times New Roman'
        
        # Create the field XML with proper namespace
        fld_xml = '''
            <w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                        xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" 
                        w:instr="PAGE"/>
        '''
        page_number._element.append(docx.oxml.parse_xml(fld_xml))
    
    # Save modified document
    output_path = f"{name}_{roll_no}.docx"
    doc.save(output_path)
    return output_path

async def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert DOCX file to PDF using LibreOffice."""
    try:
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Use LibreOffice to convert DOCX to PDF
        cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(docx_path), docx_path]
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            raise Exception(f"Conversion failed: {stderr.decode()}")
        
        return pdf_path
        
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {str(e)}")
        raise

def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Log the error and send a message to the user."""
    logger.error(f"Exception while handling an update: {context.error}")

    if isinstance(context.error, Conflict):
        logger.error("Bot instance conflict detected. Please ensure only one instance is running.")
        return

    if isinstance(context.error, (TimedOut, NetworkError)):
        logger.error(f"Network error occurred: {context.error}")
        return

    if update and isinstance(update, Update) and update.effective_message:
        error_message = "❌ An error occurred while processing your request. Please try again."
        try:
            update.effective_message.reply_text(error_message)
        except Exception as e:
            logger.error(f"Error sending error message: {e}")

def main() -> None:
    """Start the bot."""
    global application
    try:
        # Create the Application with a clean shutdown
        token = os.getenv('TELEGRAM_BOT_TOKEN')
        if not token:
            logger.error("No token found. Please set TELEGRAM_BOT_TOKEN in .env file")
            return

        # Create the Application with persistence
        application = (
            Application.builder()
            .token(token)
            .concurrent_updates(True)
            .build()
        )

        # Add conversation handler
        conv_handler = ConversationHandler(
            entry_points=[
                CommandHandler('start', start),
                MessageHandler(filters.Document.ALL, handle_document)
            ],
            states={
                WAITING_FOR_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_name)],
                WAITING_FOR_ROLLNO: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_rollno)]
            },
            fallbacks=[CommandHandler('cancel', cancel)],
            name="main_conversation",
            persistent=False
        )

        # Add handlers
        application.add_handler(conv_handler)
        
        # Add error handler
        application.add_error_handler(error_handler)

        # Run the bot in polling mode
        application.run_polling()

    except Exception as e:
        logger.error(f"Error starting bot: {e}")
        asyncio.run(cleanup())
        raise

if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        logger.info("Bot stopped by user")
    except Exception as e:
        logger.error(f"Bot stopped due to error: {e}")
    finally:
        asyncio.run(cleanup()) 